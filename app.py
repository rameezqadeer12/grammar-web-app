import os, re, json, logging, tempfile
from typing import List, Dict, Any, Tuple

import requests
from flask import Flask, render_template, request, send_file, abort, Response, jsonify
from flask_cors import CORS
from dotenv import load_dotenv
from markupsafe import escape
from docx import Document

# Rate limiting
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address

# Caching
from cachetools import TTLCache



load_dotenv()

GROQ_API_KEY = os.getenv("GROQ_API_KEY", "").strip()
APP_API_KEY  = os.getenv("APP_API_KEY", "").strip()   # optional
LT_API_URL   = os.getenv("LT_API_URL", "https://api.languagetool.org/v2/check").strip()

MAX_TEXT_CHARS = int(os.getenv("MAX_TEXT_CHARS", "12000"))
MAX_FILE_MB    = int(os.getenv("MAX_FILE_MB", "3"))
MAX_FILE_BYTES = MAX_FILE_MB * 1024 * 1024

# -----------------------------
# 2) LOGGING
# -----------------------------
logging.basicConfig(
    level=os.getenv("LOG_LEVEL", "INFO").upper(),
    format="%(asctime)s | %(levelname)s | %(name)s | %(message)s"
)
log = logging.getLogger("grammar-app")

# -----------------------------
# 3) GROQ INIT
# -----------------------------
groq_client = None
if GROQ_API_KEY:
    try:
        from groq import Groq
        groq_client = Groq(api_key=GROQ_API_KEY)
        log.info("Groq Loaded Successfully")
    except Exception as e:
        log.exception("Groq Init Error: %s", e)
else:
    log.warning("GROQ_API_KEY missing/empty: Groq features disabled")

# -----------------------------
# 4) FLASK APP
# -----------------------------
app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = MAX_FILE_BYTES
CORS(app)

limiter = Limiter(
    get_remote_address,
    app=app,
    default_limits=["60 per minute"]
)

# -----------------------------
# 5) CACHES
# -----------------------------
lt_cache   = TTLCache(maxsize=2000, ttl=60 * 30)   # 30 min
groq_cache = TTLCache(maxsize=2000, ttl=60 * 60)   # 60 min
gen_cache  = TTLCache(maxsize=200, ttl=60 * 30)

http = requests.Session()

# -----------------------------
# 6) BLACKLAW LOAD
# -----------------------------
try:
    with open("blacklaw_terms.json", "r", encoding="utf-8") as f:
        BLACKLAW = json.load(f)
except Exception:
    BLACKLAW = {}
    log.warning("blacklaw_terms.json not found/invalid. Meanings disabled.")

def normalize_key(word: str) -> str:
    return re.sub(r"[^a-z\s]", "", word.lower().strip())

# -----------------------------
# 7) LEGAL FIXES
# -----------------------------
LEGAL_FIX = {
    "suo moto": "suo motu",
    "prima facia": "prima facie",
    "mens reaa": "mens rea",
    "ratio decedendi": "ratio decidendi",
    "audi alteram partum": "audi alteram partem",
}

# ✅ CHANGE: Do NOT ignore any words (tenses words included)
IGNORE_WORDS = set()

def is_reference_like(line: str) -> bool:
    s = line.strip()
    return (
        not s
        or "http" in s
        or "www." in s
        or "doi" in s.lower()
        or re.match(r"^\[\d+\]$", s)
        or re.match(r"^\(.+\d{4}.*\)$", s)
    )

# -----------------------------
# 8) SIMPLE AUTH (optional)
# -----------------------------
def require_api_key():
    """
    Optional auth:
    - Header: X-API-Key
    - OR form field: api_key
    If APP_API_KEY is empty => auth disabled.
    """
    if not APP_API_KEY:
        return
    provided = (request.headers.get("X-API-Key", "") or "").strip()
    if not provided:
        provided = (request.form.get("api_key", "") or "").strip()
    if not provided or provided != APP_API_KEY:
        abort(401, description="Unauthorized: missing/invalid API key")

# -----------------------------
# 9) LANGUAGETOOL
# -----------------------------
def lt_check_sentence(sentence: str) -> Dict[str, Any]:
    sentence = sentence.strip()
    if not sentence:
        return {"matches": []}

    if sentence in lt_cache:
        return lt_cache[sentence]

    try:
        data = {"text": sentence, "language": "en-US"}
        r = http.post(LT_API_URL, data=data, timeout=10)
        r.raise_for_status()
        out = r.json()
    except Exception as e:
        log.warning("LT error: %s", e)
        out = {"matches": []}

    lt_cache[sentence] = out
    return out

# -----------------------------
# 10) LEGAL DETECTOR
# -----------------------------
def detect_legal(sentence: str) -> List[Tuple[str, str, str]]:
    results = []
    for wrong, correct in LEGAL_FIX.items():
        if re.search(rf"\b{re.escape(wrong)}\b", sentence, re.IGNORECASE):
            meaning = BLACKLAW.get(normalize_key(correct), "")
            results.append((wrong, correct, meaning))
    return results

# -----------------------------
# 11) GROQ WORD-ONLY CHECK (safe JSON)
# -----------------------------
def groq_word_check(sentence: str, lt_wrong_words: List[str]) -> List[Dict[str, str]]:
    if (not groq_client) or (not lt_wrong_words):
        return []

    # ✅ CHANGE: No ignore filtering; only de-duplicate
    lt_wrong_words = list(dict.fromkeys([w.strip() for w in lt_wrong_words if w.strip()]))
    if not lt_wrong_words:
        return []

    cache_key = "WORD||" + sentence + "||" + "|".join(sorted([w.lower() for w in lt_wrong_words]))
    if cache_key in groq_cache:
        return groq_cache[cache_key]

    prompt = f"""
Only correct these words (do NOT rewrite the whole sentence): {lt_wrong_words}

Apply these exact legal fixes if present:
- suo moto → suo motu
- prima facia → prima facie
- mens reaa → mens rea
- ratio decedendi → ratio decidendi
- audi alteram partum → audi alteram partem

Output ONLY valid JSON array exactly like:
[{{"wrong":"old","suggestion":"new"}}]

Sentence: "{sentence}"
""".strip()

    try:
        response = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            max_tokens=350
        )
        raw = response.choices[0].message.content.strip()

        m = re.search(r"\[(?:.|\n)*\]", raw)
        if not m:
            groq_cache[cache_key] = []
            return []

        json_str = re.sub(r",\s*]", "]", m.group(0))
        arr = json.loads(json_str)
        if not isinstance(arr, list):
            arr = []

        cleaned = []
        for item in arr:
            w = str(item.get("wrong", "")).strip()
            s = str(item.get("suggestion", "")).strip()
            if w and s and w.lower() !=s.lower():
                cleaned.append({"wrong": w, "suggestion": s})

        groq_cache[cache_key] = cleaned
        return cleaned

    except Exception as e:
        log.warning("Groq word-check error: %s", e)
        return []

# -----------------------------
# 12) GROQ FULL REWRITE (tenses+punc+grammar+spelling)
# -----------------------------
def groq_rewrite_sentence(sentence: str) -> str:
    """
    Full rewrite => tenses, punctuation, grammar, spelling correction.
    Keep meaning same.
    """
    if not groq_client:
        return sentence

    s = sentence.strip()
    if not s:
        return s

    cache_key = "REWRITE||" + s
    if cache_key in groq_cache:
        return groq_cache[cache_key]

    prompt = f"""
Rewrite this sentence in correct, clear English.
Fix: tenses, punctuation, grammar, spelling, subject-verb agreement.
Keep the meaning the same (do not add new facts).
Keep it as ONE sentence (do not split into multiple sentences).

Also fix these legal phrases if they appear:
- suo moto → suo motu
- prima facia → prima facie
- mens reaa → mens rea
- ratio decedendi → ratio decidendi
- audi alteram partum → audi alteram partem

Output ONLY the corrected sentence. No quotes. No extra text.

Sentence: {s}
""".strip()

    try:
        response = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            max_tokens=220
        )
        out = (response.choices[0].message.content or "").strip()
        if not out:
            out = s
        groq_cache[cache_key] = out
        return out
    except Exception as e:
        log.warning("Groq rewrite error: %s", e)
        return s

# -----------------------------
# 13) GENERATE LONG PARAGRAPH WITH MISTAKES
# -----------------------------
def generate_mistake_paragraph(topic: str = "court case") -> str:
    """
    Generates a long paragraph intentionally filled with mistakes:
    tenses, punctuation, grammar, spelling, etc.
    """
    topic = (topic or "court case").strip()
    cache_key = topic.lower()
    if cache_key in gen_cache:
        return gen_cache[cache_key]

    # fallback (if Groq missing)
    if not groq_client:
        text = (
            "The judge take suo motu action without properly hearing the defendant and he dont give fair "
            "oppurtunity for defence, which are against natural justice. The court say prima facie evidences "
            "was enough but ratio decidendi were not clear, and mens rea is not establish, however it ignore. "
            "Witnesses statement was contradictory, some witness change there version in cross examination and "
            "others contradict each other, but judgement didnt considered it properly, so it seems court rely on "
            "assumptions then law, and burdens of proofs shifted on accused which is not correct, moreover "
            "procedural irregularities was overlooked and precedent case laws were not analyse, making decision "
            "arbitrary capricious and not sustainable in eyes of law."
        )
        gen_cache[cache_key] = text
        return text

    prompt = f"""
Write ONE long paragraph (minimum 180-220 words) about: "{topic}".

IMPORTANT: Make it intentionally incorrect English with MANY mistakes:
- tense mistakes
- punctuation mistakes
- grammar mistakes
- spelling mistakes
- subject-verb agreement mistakes
- wrong plural/singular usage
- awkward legal phrasing

It should still be understandable.
Output ONLY the paragraph.
""".strip()

    try:
        r = groq_client.chat.completions.create(
            model="llama-3.3-70b-versatile",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.7,
            max_tokens=450
        )
        out = (r.choices[0].message.content or "").strip()
        if not out:
            out = "The court take action but the reasoning are not clear and it make many errors."
        gen_cache[cache_key] = out
        return out
    except Exception as e:
        log.warning("Generate paragraph error: %s", e)
        return "The court take action but the reasoning are not clear and it make many errors."

# -----------------------------
# 14) BUILD HIGHLIGHTED HTML
# -----------------------------
def process_text_line_by_line(text: str, mode: str = "word") -> str:
    """
    mode:
      - 'word'   => highlight + suggestions
      - 'rewrite'=> full corrected lines returned (no highlights)
    """
    lines = text.split("\n")
    final_html = []

    for line in lines:
        if not line.strip():
            final_html.append("<p></p>")
            continue

        # XSS safe display
        safe_line = escape(line)

        if is_reference_like(line):
            final_html.append(f"<p>{safe_line}</p>")
            continue

        working = line

        # FULL REWRITE MODE (tenses + punctuation + grammar + spelling)
        if mode == "rewrite":
            corrected = groq_rewrite_sentence(working)
            final_html.append(f"<p>{escape(corrected)}</p>")
            continue

        # WORD MODE (highlight suggestions)
        html_line = str(safe_line)
        lt_res = lt_check_sentence(working)

        lt_wrong_words = []
        for m in lt_res.get("matches", []):
            wrong = working[m["offset"]:m["offset"] + m["length"]]
            # ✅ CHANGE: no ignore; keep tense words too
            if wrong.strip() and len(wrong.strip()) > 0:
                lt_wrong_words.append(wrong)

        legal_hits = detect_legal(working)
        groq_hits  = groq_word_check(working, lt_wrong_words)

        combined = {}

        for wrong, correct, meaning in legal_hits:
            key = wrong.lower()
            combined[key] = {"original": wrong, "black": correct, "groq": None, "meaning": meaning}

        for g in groq_hits:
            wrong_raw = (g.get("wrong") or "").strip()
            suggestion = (g.get("suggestion") or "").strip()
            if not wrong_raw or not suggestion:
                continue

            key = wrong_raw.lower()
            if key not in combined:
                mm = re.search(rf"\b{re.escape(wrong_raw)}\b", working, re.IGNORECASE)
                original = mm.group(0) if mm else wrong_raw
                combined[key] = {"original": original, "black": None, "groq": None, "meaning": ""}

            combined[key]["groq"] = suggestion

        for _, data in combined.items():
            original_word = data["original"]
            black = data["black"] or ""
            groq  = data["groq"] or ""
            meaning = data["meaning"] or ""

            ow = escape(original_word)
            span = (
                f"<span class='grammar-wrong' "
                f"data-wrong='{escape(original_word)}' "
                f"data-black='{escape(black)}' "
                f"data-groq='{escape(groq)}' "
                f"data-meaning='{escape(meaning)}'>"
                f"{ow}</span>"
            )

            # ✅ CHANGE: stable replace (no regex boundary issues)
            html_line = html_line.replace(str(ow), span, 1)

        final_html.append(f"<p>{html_line}</p>")

    return "\n".join(final_html)

# -----------------------------
# 15) ROUTES
# -----------------------------
@app.route("/", methods=["GET", "POST"])
@limiter.limit("20 per minute")
def index():
    if request.method == "POST":
        require_api_key()

        text_input = (request.form.get("text", "") or "").strip()
        mode = (request.form.get("mode", "word") or "word").strip().lower()
        file = request.files.get("file")

        text = text_input
        if file and file.filename.lower().endswith(".docx"):
            doc = Document(file)
            text = "\n".join([p.text for p in doc.paragraphs])

        if len(text) > MAX_TEXT_CHARS:
            abort(413, description=f"Text too long. Max {MAX_TEXT_CHARS} chars allowed.")

        output = process_text_line_by_line(text, mode=mode)
        return render_template("result.html", highlighted_html=output, mode=mode)

    return render_template("index.html")

@app.route("/check", methods=["POST"])
@limiter.limit("30 per minute")
def check_text():
    require_api_key()
    text = (request.form.get("text", "") or "")
    mode = (request.form.get("mode", "word") or "word").strip().lower()

    if len(text) > MAX_TEXT_CHARS:
        abort(413, description=f"Text too long. Max {MAX_TEXT_CHARS} chars allowed.")

    output = process_text_line_by_line(text, mode=mode)
    return Response(output, mimetype="text/html")

@app.route("/generate_sample", methods=["POST"])
@limiter.limit("15 per minute")
def generate_sample():
    require_api_key()
    topic = (request.form.get("topic", "") or "court case").strip()
    para = generate_mistake_paragraph(topic)
    return jsonify({"text": para})

@app.route("/download_corrected", methods=["POST"])
@limiter.limit("10 per minute")
def download_corrected():
    require_api_key()

    final_text = (request.form.get("final_text", "") or "")
    if len(final_text) > MAX_TEXT_CHARS:
        abort(413, description=f"Text too long. Max {MAX_TEXT_CHARS} chars allowed.")

    # In rewrite mode, final_text is already corrected plain text.
    # In word mode, we still allow replacements.
    try:
        replacements = json.loads(request.form.get("replacements", "[]"))
        if not isinstance(replacements, list):
            replacements = []
    except Exception:
        replacements = []

    doc = Document()
    for line in final_text.split("\n"):
        doc.add_paragraph(line)

    for para in doc.paragraphs:
        for rep in replacements:
            wrong = str(rep.get("old", "")).strip()
            correct = str(rep.get("new", "")).strip()
            if not wrong or not correct:
                continue
            para.text = re.sub(rf"\b{re.escape(wrong)}\b", correct, para.text, flags=re.IGNORECASE)

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_path = tmp.name
    tmp.close()

    doc.save(tmp_path)
    return send_file(tmp_path, as_attachment=True, download_name="Corrected_Final_Output.docx")

@app.route("/health", methods=["GET"])
def health():
    return {"status": "ok"}
@app.route("/check_grammar", methods=["POST"])
@limiter.limit("30 per minute")
def check_grammar():
    require_api_key()  # optional, if you want to enforce API key

    data = request.get_json(force=True) or {}
    text = (data.get("text", "") or "").strip()
    mode = (data.get("mode", "word") or "word").strip().lower()

    if len(text) > MAX_TEXT_CHARS:
        abort(413, description=f"Text too long. Max {MAX_TEXT_CHARS} chars allowed.")

    output = process_text_line_by_line(text, mode=mode)
    return jsonify({"suggestions": output})    

# DEV ONLY
if __name__ == "__main__":
    port = int(os.environ.get("PORT", "5000"))
    app.run(host="0.0.0.0", port=port, debug=False)



